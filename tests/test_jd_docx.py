"""Branded DOCX export — produces a valid .docx that mirrors the MY HOME GROUP
reference layout: a Job Title / Location / Reporting / Company Link block with
teal labels, a logo in the page header, a 'Page X of Y' footer, and the JD body
sections (the leading Markdown header table and footer rule are stripped)."""

import io

from docx import Document

from app import jd_docx

RECORD = {
    "title": "GET - PMO",
    "location": "Hyderabad",
    "reporting": "Head - PMO",
    "content": (
        "| | |\n|---|---|\n"
        "| **Job Title** | GET - PMO |\n"
        "| **Location** | Hyderabad |\n"
        "| **Reporting** | Head - PMO |\n"
        "| **Company Link** | [www.x.com](https://www.x.com) |\n\n"
        "## About the Opportunity\nGreat role with growth.\n\n"
        "## Key Responsibilities\n"
        "**1. Reporting**\n- Build MIS reports.\n- Track KPIs.\n\n"
        "## Qualifications & Skills\n**Education:** B.E. Civil\n\n"
        "---\n_Company footer._\n"
    ),
}


def test_docx_is_valid_and_structured():
    data = jd_docx.jd_docx_bytes(RECORD)
    assert data[:2] == b"PK"  # docx is a zip
    doc = Document(io.BytesIO(data))

    full = "\n".join(p.text for p in doc.paragraphs)

    # header info block (teal-label key/value lines) present
    assert "Job Title:" in full and "GET - PMO" in full
    assert "Location:" in full and "Hyderabad" in full
    assert "Reporting:" in full and "Head - PMO" in full
    assert "Company Link:" in full

    # body sections rendered, header-table markdown + footer stripped
    assert "About the Opportunity" in full
    assert "Build MIS reports." in full
    assert "| **Job Title**" not in full      # raw markdown header not leaked
    assert "---" not in full                   # footer rule stripped


def test_docx_has_page_furniture():
    """Logo header + 'Page X of Y' footer carrying the job title label."""
    doc = Document(io.BytesIO(jd_docx.jd_docx_bytes(RECORD)))
    section = doc.sections[0]

    footer_text = "\n".join(p.text for p in section.footer.paragraphs)
    assert "GET - PMO" in footer_text
    assert "Page" in footer_text

    # the PAGE / NUMPAGES fields are emitted in the footer XML
    footer_xml = section.footer.paragraphs[0]._p.xml
    assert "PAGE" in footer_xml and "NUMPAGES" in footer_xml


def test_docx_without_header_table_or_footer():
    rec = {"title": "Analyst", "content": "## Overview\nDoes analysis."}
    doc = Document(io.BytesIO(jd_docx.jd_docx_bytes(rec)))
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Does analysis." in full

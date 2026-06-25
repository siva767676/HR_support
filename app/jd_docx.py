"""Render a repository JD into a branded .docx that mirrors the MY HOME GROUP
reference layout:

  * the company logo top-right on every page (document header),
  * a Job Title / Location / Reporting / Company Link block with teal labels,
  * underlined section headings, numbered responsibility themes, bullet points,
  * a "<Job Title>  Page X of Y" footer on every page.

python-docx is already a project dependency. Drop a logo at config.COMPANY_LOGO
(defaults to the bundled MY HOME GROUP lockup) to brand the export."""

from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from . import config

# Reference palette: teal field labels, blue links, near-black body text.
_TEAL = RGBColor(0x1F, 0x6E, 0x7C)
_LINK = RGBColor(0x05, 0x63, 0xC1)
_INK = RGBColor(0x1A, 0x1A, 0x1A)
_BODY_FONT = "Calibri"


# --------------------------- low-level docx helpers ---------------------------

def _add_inline(paragraph, text: str) -> None:
    """Add text to a paragraph, honouring **bold** segments."""
    for i, seg in enumerate(re.split(r"\*\*", text)):
        if seg == "":
            continue
        run = paragraph.add_run(seg)
        run.bold = i % 2 == 1  # odd segments were between ** **


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Append a real, clickable hyperlink (blue, underlined) to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _add_field(paragraph, instruction: str) -> None:
    """Append a Word field (e.g. PAGE / NUMPAGES) to a paragraph."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


# ------------------------------ content rendering -----------------------------

def _strip_header_and_footer(md: str) -> str:
    """Drop the leading Markdown header table and trailing '---' footer, leaving
    just the JD section body (rendered natively below)."""
    lines = md.splitlines()
    i = 0
    while i < len(lines) and (lines[i].strip().startswith("|") or not lines[i].strip()):
        i += 1
    rest = lines[i:]
    for j in range(len(rest) - 1, -1, -1):
        if rest[j].strip() == "---":
            rest = rest[:j]
            break
    return "\n".join(rest).strip()


def _heading(doc, text: str, *, size: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(size)
    run.font.color.rgb = _INK


def _render_body(doc, body: str) -> None:
    theme_re = re.compile(r"^\*\*\d+\.")  # "**1. Theme**" responsibility groups
    for raw in body.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            _heading(doc, line[4:].strip(), size=11)
        elif line.startswith("## "):
            _heading(doc, line[3:].strip(), size=13)
        elif line.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line.lstrip()[2:])
        elif theme_re.match(line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            _add_inline(p, line.strip())
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)


# -------------------------------- page furniture ------------------------------

def _build_header(section) -> None:
    """Logo, right-aligned, repeated on every page."""
    logo = Path(config.COMPANY_LOGO)
    para = section.header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if logo.exists():
        try:
            para.add_run().add_picture(str(logo), width=Inches(1.25))
        except Exception:  # noqa: BLE001 - a bad image must not break the export
            pass


def _build_footer(section, label: str) -> None:
    """'<Job Title>  Page X of Y', right-aligned, every page."""
    para = section.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(f"{label}    Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
    _add_field(para, "PAGE")
    mid = para.add_run(" of ")
    mid.font.size = Pt(9)
    mid.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)
    _add_field(para, "NUMPAGES")


def _kv(doc, label: str, value: str, *, link: str | None = None) -> None:
    """A 'Label:    Value' line: teal bold label, tab, bold black value (or link)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(1.55), WD_TAB_ALIGNMENT.LEFT)
    lab = p.add_run(label)
    lab.bold = True
    lab.font.color.rgb = _TEAL
    p.add_run("\t")
    if link:
        p.add_run("(")
        _add_hyperlink(p, link, value)
        p.add_run(")")
    else:
        val = p.add_run(value)
        val.bold = True
        val.font.color.rgb = _INK


# ----------------------------------- export -----------------------------------

def jd_docx_bytes(record: dict) -> bytes:
    doc = Document()

    # Base typography to match the reference (clean Office sans).
    normal = doc.styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _INK

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    title = record.get("title", "") or "Job Description"
    _build_header(section)
    _build_footer(section, title)

    # Header info block (teal labels), mirroring the reference document.
    _kv(doc, "Job Title:", title)
    _kv(doc, "Location:", record.get("location") or "—")
    _kv(doc, "Reporting:", record.get("reporting") or "—")
    site = config.COMPANY_WEBSITE
    _kv(doc, "Company Link:", site, link=f"https://{site}")
    doc.add_paragraph()

    _render_body(doc, _strip_header_and_footer(record["content"]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

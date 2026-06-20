import io
import os
import re


def extract_text(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        text = _from_pdf(data)
    elif ext == ".docx":
        text = _from_docx(data)
    elif ext in (".txt", ".md"):
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type '{ext}' — use PDF, DOCX or TXT")

    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        raise ValueError(f"No text could be extracted from {filename}")
    return text


def _from_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF

    with fitz.open(stream=data, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def _from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
